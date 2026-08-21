"""
build_wireframe_doc.py
Compiles all 10 SavePlate screenshots into a formatted Word document.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SAVE_DIR = "/Users/shekharbhatt/Desktop/saveplate"
OUTPUT_FILE = os.path.join(SAVE_DIR, "SavePlate_Wireframes.docx")

# (section_number, title, figure_number, figure_caption, screenshot_filename)
SCREENS = [
    (
        "5.1",
        "UC1 — Register / Create Account Page",
        5,
        "Register Page Wireframe — SavePlate Registration Screen",
        "screen_01_register.png",
    ),
    (
        "5.2",
        "UC1 — 2FA Setup Page",
        6,
        "2FA Wireframe — Two-Factor Authentication Verification Screen",
        "screen_02_2fa.png",
    ),
    (
        "5.3",
        "UC1 — Account Privacy and Security Settings Page",
        7,
        "Account and Privacy Settings Wireframe — Settings Management Screen",
        "screen_03_settings.png",
    ),
    (
        "5.4",
        "UC2 — Food Inventory Page",
        8,
        "Food Inventory Wireframe — Inventory Management Screen",
        "screen_04_inventory.png",
    ),
    (
        "5.5",
        "UC2 — Add Food Item Dialog",
        9,
        "Add Food Item Wireframe — Add Food Modal Dialog Screen",
        "screen_05_add_food.png",
    ),
    (
        "5.6",
        "UC3 — Browse and Donate Page",
        10,
        "Browse and Donate Wireframe — Community Donations Board Screen",
        "screen_06_browse.png",
    ),
    (
        "5.7",
        "UC4 — Food Analytics Dashboard",
        11,
        "Food Analytics Wireframe — Analytics and Metrics Dashboard Screen",
        "screen_07_analytics.png",
    ),
    (
        "5.8",
        "UC5 — Notifications Page",
        12,
        "Notifications Wireframe — Notification and Alerts Center Screen",
        "screen_08_notifications.png",
    ),
    (
        "5.9",
        "UC6 — Weekly Meal Planner",
        13,
        "Weekly Meal Planner Wireframe — 7-Day Meal Planning Calendar Screen",
        "screen_10_mealplan.png",
    ),
    (
        "5.10",
        "Dashboard Overview",
        14,
        "Dashboard Wireframe — Main Dashboard Overview Screen",
        "screen_09_dashboard.png",
    ),
]


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading("SavePlate", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    sub = doc.add_paragraph("UI Wireframes & Screen Designs")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    desc = doc.add_paragraph(
        "This document contains full-resolution screenshots of all key application screens, "
        "corresponding to Use Cases UC1 – UC6 as defined in the SavePlate requirements specification."
    )
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.runs[0].font.size = Pt(10)
    desc.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    add_horizontal_rule(doc)
    doc.add_paragraph()

    # Screens
    for section_num, title_text, figure_num, caption_text, filename in SCREENS:
        img_path = os.path.join(SAVE_DIR, filename)

        # Section heading
        heading = doc.add_heading("", level=2)
        heading.clear()
        run = heading.add_run(f"{section_num}  {title_text}")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1A, 0x5E, 0xA3)
        heading.paragraph_format.space_before = Pt(16)
        heading.paragraph_format.space_after = Pt(8)

        if os.path.exists(img_path):
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = img_para.add_run()
            run_img.add_picture(img_path, width=Inches(5.8))
            print(f"  ✓ Embedded: {filename}")
        else:
            placeholder = doc.add_paragraph(f"[ Screenshot not found: {filename} ]")
            placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            placeholder.runs[0].italic = True
            placeholder.runs[0].font.color.rgb = RGBColor(0xAA, 0x00, 0x00)
            print(f"  ✗ Missing: {filename}")

        # Figure caption
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(f"Figure {figure_num}: {caption_text}")
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        caption.paragraph_format.space_after = Pt(6)

        add_horizontal_rule(doc)

    # Footer
    doc.add_paragraph()
    note = doc.add_paragraph("SavePlate — Smart Food Waste Reduction App | Screen Designs (Sections 5.1 – 5.10)")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].font.size = Pt(8)
    note.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(OUTPUT_FILE)
    print(f"\n✓ Word document saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_doc()
